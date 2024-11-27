import math
import os
from io import BytesIO

#тут зашитые значения
from settings import *

from flask import Flask, render_template, request, jsonify, send_file

#библиотеки для работы с экселем
import pandas as pd
import openpyxl

#библиотеки для работы с пдф
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.pagesizes import letter, landscape
from reportlab.pdfgen import canvas
from reportlab.platypus import Table, TableStyle, SimpleDocTemplate
from reportlab.lib import colors

#библиотека для работы с вордом
from docx import Document

app = Flask(__name__)
app.secret_key = os.urandom(24) #ключ защищающий сессию пользователя
@app.route('/')
def index():
    return render_template('index.html')

#функция, подготавлиявающая полученные расчеты для экспорта
def format_dicts_side_by_side(dicts):
    """Форматирует словари, подгоняя их под ключи 180h, 168h, 79h, 180h_pr, 180h_night"""

    column_keys = ['180h', '168h', '79h', '180h_pr', '180h_night']

    table = {key: [] for key in column_keys}

    for d in dicts:
        for key, value in d.items():
            if key in column_keys:
                table[key].append(value)
            elif key == '180-79h':
                table['180h'].append(value)
                table['168h'].append(value)
                table['79h'].append(value)
            elif key == '180h_pr_night':
                table['180h_pr'].append(value)
                table['180h_night'].append(value)
            else:
                for k in column_keys:
                    table[k].append('')
    return table

#факт расчет
def fact(total_files, day_files, night_files, day_pr_files, machines_180h, machines_168h, machines_79h, machines_180h_night):
    # факт среднее количество файлов в месяц для 180-79ч 180ч пр/вых 180чночь
    fact_mid_files_month = {
        '180-79h': day_files,
        '180h_pr': day_pr_files,
        '180h_night': night_files
    }

    # факт кол-во машин для 180ч, 168ч, 79ч, 180ч пр/вых, 180ч ночь соответсвтенно
    fact_mashine_count = {
        '180h': machines_180h,
        '168h': machines_168h,
        '79h': machines_79h,
        '180h_pr': machines_180h_night,  # надо уточнить берем ли мы тут значение из F17(см калькулятор.xls)
        '180h_night': machines_180h_night
    }

    # факт максимальное количество файлов для 180ч, 168ч, 79ч, 180ч пр/вых, 180ч ночь соответсвтенно
    fact_max_files = {
        '180h': round(fact_mashine_count['180h'] * max_files_month['180h']),
        '168h': round(fact_mashine_count['168h'] * max_files_month['168h']),
        '79h': round(fact_mashine_count['79h'] * max_files_month['79h']),
        '180h_pr': round(fact_mashine_count['180h_pr'] * max_files_month['180h_pr']),
        '180h_night': round(fact_mashine_count['180h_night'] * max_files_month['180h_night'])
    }

    # факт разница нагрузки
    fact_diff_stress = {
        '180-79h': round(
            fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h'] - fact_mid_files_month['180-79h']),
        '180h_pr': round(fact_max_files['180h_pr'] - fact_mid_files_month['180h_pr']),
        '180h_night': round(fact_max_files['180h_night'] - fact_mid_files_month['180h_night'])
    }

    # факт нагрузка в %
    fact_percent_stress = {
        '180-79h': round(100 * fact_mid_files_month['180-79h'] / (
                    fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h'])),
        '180h_pr': round(100 * fact_mid_files_month['180h_pr'] / fact_max_files['180h_pr']),
        '180h_night': round(100 * fact_mid_files_month['180h_night'] / fact_max_files['180h_night'])
    }

    # факт нехватка машин
    fact_mashine_loss = {
        '180h': 'black',
        '168h': 0,
        '79h': -1,
        '180h_pr_night': 0
    }
    if fact_percent_stress['180-79h'] < 86:
        fact_mashine_loss['168h'] = 0
    else:
         fact_mashine_loss['168h'] = round( ((fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h'])  * (fact_percent_stress['180-79h'] - 86)) / 86 / max_files_month['168h'])

    if fact_mid_files_month['180-79h'] / (fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h'] + fact_mashine_loss['168h'] * max_files_month['168h']) < 0.86:
        fact_mashine_loss['79h'] = 0
    else:
        fact_mashine_loss['79h'] = math.ceil((fact_mid_files_month['180-79h'] / (fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h'] + fact_mashine_loss['168h'] * max_files_month['168h']) - 0.86) * (fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h'] + fact_mashine_loss['168h'] * max_files_month['168h']) / 86 / max_files_month['79h'])

    if fact_diff_stress['180h_night'] < 0:
        fact_mashine_loss['180h_pr_night'] = round(fact_max_files['180h_night'] / max_files_month['180h_night'] * 1.8)
    elif round((fact_percent_stress['180h_night'] + fact_percent_stress['180h_pr']) / 2) < 80:
        fact_mashine_loss['180h_pr_night'] = 0
    else:
        fact_mashine_loss['180h_pr_night'] = round(max_files_month['180h_night'] / fact_max_files['180h_night'])

    fact_mashine_loss['180h'] = fact_mashine_loss['180h_pr_night']

    data[0] = format_dicts_side_by_side([fact_mid_files_month, fact_mashine_count, fact_max_files, fact_diff_stress, fact_percent_stress, fact_mashine_loss])

    fact_result = (
        #"Факт: <br><br>"
        #f"{format_dicts_side_by_side([fact_mid_files_month, fact_mashine_count, fact_max_files, fact_diff_stress, fact_percent_stress, fact_mashine_loss])}<br><br>"
        f"Фактическая нехватка машин<br>"
        f"для 168ч: {fact_mashine_loss['168h']}<br>"
        f"для 79ч: {fact_mashine_loss['79h']}<br>"
        f"для 180ч пр/ночь: {fact_mashine_loss['180h_pr_night']}<br><br>"
    )

    return fact_result


#план расчет
def plan(total_files, day_files, night_files, day_pr_files, machines_180h, machines_168h, machines_79h, machines_180h_night, new_users):
    # факт среднее количество файлов в месяц для 180-79ч 180ч пр/вых 180чночь
    fact_mid_files_month = {
        '180-79h': day_files,
        '180h_pr': day_pr_files,
        '180h_night': night_files
    }

    # факт кол-во машин для 180ч, 168ч, 79ч, 180ч пр/вых, 180ч ночь соответсвтенно
    fact_mashine_count = {
        '180h': machines_180h,
        '168h': machines_168h,
        '79h': machines_79h,
        '180h_pr': machines_180h_night,  # надо уточнить берем ли мы тут значение из F17(см калькулятор.xls)
        '180h_night': machines_180h_night
    }

    # факт максимальное количество файлов для 180ч, 168ч, 79ч, 180ч пр/вых, 180ч ночь соответсвтенно
    fact_max_files = {
        '180h': round(fact_mashine_count['180h'] * max_files_month['180h']),
        '168h': round(fact_mashine_count['168h'] * max_files_month['168h']),
        '79h': round(fact_mashine_count['79h'] * max_files_month['79h']),
        '180h_pr': round(fact_mashine_count['180h_pr'] * max_files_month['180h_pr']),
        '180h_night': round(fact_mashine_count['180h_night'] * max_files_month['180h_night'])
    }

    # факт разница нагрузки
    fact_diff_stress = {
        '180-79h': round(
            fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h'] - fact_mid_files_month['180-79h']),
        '180h_pr': round(fact_max_files['180h_pr'] - fact_mid_files_month['180h_pr']),
        '180h_night': round(fact_max_files['180h_night'] - fact_mid_files_month['180h_night'])
    }

    # факт нагрузка в %
    fact_percent_stress = {
        '180-79h': round(100 * fact_mid_files_month['180-79h'] / (
                    fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h'])),
        '180h_pr': round(100 * fact_mid_files_month['180h_pr'] / fact_max_files['180h_pr']),
        '180h_night': round(100 * fact_mid_files_month['180h_night'] / fact_max_files['180h_night'])
    }

    # факт нехватка машин
    fact_mashine_loss = {
        '180h': 'black',
        '168h': 0,
        '79h': -1,
        '180h_pr_night': 0
    }

    # среднее количесво файлов новых уз в месяц
    plan_mid_fileUZ_month = {
        '180-79h': 'задастся ниже',
        '180h_pr': round((new_users * 1.3) * 0.15),
        '180h_night': round((new_users * 1.3) * 0.27)
    }
    plan_mid_fileUZ_month['180-79h'] = round(
        new_users * 1.3 - plan_mid_fileUZ_month['180h_night'] - plan_mid_fileUZ_month['180h_pr'])

    # среднее кол-во файлов с учетом новых уз в месяц
    plan_mid_newfiles_month = {
        '180-79h': round(fact_mid_files_month['180-79h'] + plan_mid_fileUZ_month['180-79h']),
        '180h_pr': round(fact_mid_files_month['180h_pr'] + plan_mid_fileUZ_month['180h_pr']),
        '180h_night': round(fact_mid_files_month['180h_night'] + plan_mid_fileUZ_month['180h_night'])
    }

    # планируемая разница нагрузки
    plan_diff_stress = {
        '180-79h': round(
            fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h'] - plan_mid_newfiles_month[
                '180-79h']),
        '180h_pr': round(fact_max_files['180h_pr'] - plan_mid_newfiles_month['180h_pr']),
        '180h_night': round(fact_max_files['180h_night'] - plan_mid_newfiles_month['180h_night'])
    }

    # планируемая разница нагрузки в процентах
    plan_percent_stress = {
        '180-79h': round(100 * plan_mid_newfiles_month['180-79h'] / (
                    fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h'])),
        '180h_pr': round(100 * plan_mid_newfiles_month['180h_pr'] / fact_max_files['180h_pr']),
        '180h_night': round(100 * plan_mid_newfiles_month['180h_night'] / fact_max_files['180h_night'])
    }

    # планируемая нехватка машин
    plan_mashine_loss = {
        '180h': 'black',
        '168h': -1,
        '79h': -2,
        '180h_pr_night': 0
    }

    if plan_percent_stress['180-79h'] < 86:
        plan_mashine_loss['168h'] = 0
    else:
        plan_mashine_loss['168h'] = round(((fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h']) * (plan_percent_stress['180-79h'] - 86) ) / 86 / max_files_month['168h'])

    #=ЕСЛИ((M4/(СУММ(R4:R6) + U5 * F5))<0,86;0;ОКРУГЛВВЕРХ((M4/(СУММ(R4:R6) + U5*F5) - 0,86)*(СУММ(R4:R6) + U5 * F5) / 0,86 / F6; 0))
    if (fact_mid_files_month['180-79h'] / ((fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h']) + plan_mashine_loss['168h'] * max_files_month['168h'])) < 86:
        plan_mashine_loss['79h'] = 0
    else:
        plan_mashine_loss['79h'] = math.ceil(((fact_mid_files_month['180-79h'] / ((fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h']) + plan_mashine_loss['168h'] * max_files_month['168h'])) - 86) * ((fact_max_files['180h'] + fact_max_files['168h'] + fact_max_files['79h']) + plan_mashine_loss['168h'] * max_files_month['168h']) / 86 / max_files_month['79h'])

    if plan_diff_stress['180h_night'] < 0:
        plan_mashine_loss['180h_pr_night'] = round(fact_max_files['180h_night'] / max_files_month['180h_night'] * 1.8)
    elif (plan_percent_stress['180h_night'] + plan_percent_stress['180h_pr']) / 2 < 80:
        plan_mashine_loss['180h_pr_night'] = 0
    else:
        plan_mashine_loss['180h_pr_night'] = round(max_files_month['180h_night'] / fact_max_files['180h_night'])

    plan_mashine_loss['180h'] = plan_mashine_loss['180h_pr_night']

    data[1] = format_dicts_side_by_side([fact_mid_files_month, {'180h': new_users, '168h': new_users, '79h': new_users,'180h_pr': new_users,'180h_night': new_users}, plan_mid_fileUZ_month, plan_mid_newfiles_month, fact_mashine_count, fact_max_files, plan_diff_stress, plan_percent_stress, plan_mashine_loss])

    plan_result = (
        #"План: <br><br>"
        #f"{format_dicts_side_by_side([fact_mid_files_month, {'180h': new_users, '168h': new_users, '79h': new_users,'180h_pr': new_users,'180h_night': new_users}, plan_mid_fileUZ_month, plan_mid_newfiles_month, fact_mashine_count, fact_max_files, plan_diff_stress, plan_percent_stress, plan_mashine_loss])}<br>"
        f"Планируемая нехватка машин<br>"
        f"для 168ч: {plan_mashine_loss['168h']}<br>"
        f"для 79ч: {plan_mashine_loss['79h']}<br>"
        f"для 180ч пр/ночь: {plan_mashine_loss['180h_pr_night']}<br>"
    )
    return plan_result


#тут хранятся данные для экспорта, полученные после расчета
data = [{}, {}]

#экспорт результатов в табличку эксель
@app.route('/export_excel')
def export_excel():
    try:
        if data[0] != {}:
            # Получаем ключи и определяем количество колонок
            keys = list(data[0].keys())
            num_columns_fact = 6
            columns_fact = [
                    'Факт среднее\n кол-во файлов\n в месяц',
                    'Факт  \nкол-во машин',
                    'Факт \nмаксимальное\n кол-во файлов',
                    'Факт разница\n нагрузки',
                    'Факт \nнагрузка в %',
                    'Факт \nнехватка машин'
                    ]
            num_columns_fact = 9
            columns_plan = [
                    'Факт \nсреднее кол-во\n файлов в месяц',
                    'Кол-во \nновых УЗ',
                    'Среднее кол-во\n файлов новых\n УЗ в месяц',
                    'Среднее  кол-во \nфайлов с учетом \nновых УЗ в месяц',
                    'Факт  \nкол-во машин',
                    'Факт \nмаксимальное \nкол-во файлов',
                    'Планируемая\nразница нагрузки',
                    'Планируемая\nнагрузка в %',
                    'Планируемая\nнехватка машин'
                    ]

            rows_fact = []
            for key in keys:
                row = [key]
                values = data[0].get(key, [])
                row.extend(values)
                rows_fact.append(row)

            headers_fact = ["ФАКТ"] + columns_fact
            df_fact = pd.DataFrame(rows_fact, columns=headers_fact)

            rows_plan = []
            if data[1] != {}:
                for key in keys:
                    row = [key]
                    values = data[1].get(key, [])
                    row.extend(values)
                    rows_plan.append(row)

            headers_plan = ["ПЛАН"] + columns_plan
            df_plan = pd.DataFrame(rows_plan, columns=headers_plan)

            output = BytesIO()
            with pd.ExcelWriter(output, engine='openpyxl') as writer:
                # Записываем первую таблицу
                df_fact.to_excel(writer, index=False, sheet_name='Sheet1', startrow=0)
                # Вычисляем начальную строку для второй таблицы с учетом заголовков и данных первой таблицы
                startrow_plan = len(df_fact) + 2  # +1 для отступа и +1 для заголовков
                # Записываем вторую таблицу
                df_plan.to_excel(writer, index=False, sheet_name='Sheet1', startrow=startrow_plan)

                # Получаем объект worksheet
                worksheet = writer.sheets['Sheet1']

                # бъединяем ячейки
                #fact
                worksheet.merge_cells(start_row=2, start_column=2, end_row=4, end_column=2)
                worksheet["B2"] = data[0]['180h'][0]

                worksheet.merge_cells(start_row=2, start_column=5, end_row=4, end_column=5)
                worksheet["E2"] = data[0]['180h'][3]

                worksheet.merge_cells(start_row=2, start_column=6, end_row=4, end_column=6)
                worksheet["F2"] = data[0]['180h'][4]

                worksheet.merge_cells(start_row=5, start_column=7, end_row=6, end_column=7)
                worksheet["G5"] = data[0]['180h_night'][5]

                #plan
                if data[1] != {}:
                    worksheet.merge_cells(start_row=9, start_column=2, end_row=11, end_column=2)
                    worksheet["B9"] = data[1]['180h'][0]

                    worksheet.merge_cells(start_row=9, start_column=3, end_row=13, end_column=3)
                    worksheet["C9"] = data[1]['180h'][1]

                    worksheet.merge_cells(start_row=9, start_column=4, end_row=11, end_column=4)
                    worksheet["D9"] = data[1]['180h'][2]

                    worksheet.merge_cells(start_row=9, start_column=5, end_row=11, end_column=5)
                    worksheet["E9"] = data[1]['180h'][3]

                    worksheet.merge_cells(start_row=9, start_column=8, end_row=11, end_column=8)
                    worksheet["H9"] = data[1]['180h'][6]

                    worksheet.merge_cells(start_row=9, start_column=9, end_row=11, end_column=9)
                    worksheet["I9"] = data[1]['180h'][7]

                    worksheet.merge_cells(start_row=12, start_column=10, end_row=13, end_column=10)
                    worksheet["J12"] = data[1]['180h_night'][8]

                # Установка ширины колонок
                for col_num, column_cells in enumerate(zip(*df_fact.values.tolist(), *df_plan.values.tolist()), 1):
                    max_length = max(
                        len(str(column_cells[0])),
                        *(len(str(cell)) for cell in column_cells if cell is not None)  # Игнорируем None
                    )
                    # Настраиваем ширину колонок с небольшим запасом
                    worksheet.column_dimensions[openpyxl.utils.get_column_letter(col_num)].width = max_length + 2

                # Установка высоты строк для первой таблицы
                for row_num, row in enumerate([df_fact.columns.tolist()] + df_fact.values.tolist(),
                                              1):  # Нумерация строк начинается с 1
                    max_height = max(
                        len(str(cell).split('\n')) for cell in row if cell is not None)  # Учитываем переносы строк
                    worksheet.row_dimensions[row_num].height = max_height * 15  # Высота строки ~15 пикселей на линию

                # Установка высоты строк для второй таблицы
                for row_num, row in enumerate([df_plan.columns.tolist()] + df_plan.values.tolist(), startrow_plan + 1):
                    max_height = max(len(str(cell).split('\n')) for cell in row if cell is not None)
                    worksheet.row_dimensions[row_num].height = max_height * 15

            output.seek(0)

            # Передача файла напрямую в ответе, не используя сессии
            return send_file(output, as_attachment=True, download_name="report.xlsx",
                             mimetype="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        else:
            return '', 204
    except Exception as ex:
        print(ex)
        return '', 204

#экспорт результатов в файл пдф
#выбор шрифта для файла пфд(локально)
pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
#выбор шрифта для файла пфд(путь для сервера)
#pdfmetrics.registerFont(TTFont('Arial', '/home/dm3tr0/mysite/fonts/arial.ttf'))
@app.route('/export_pdf')
def export_pdf():
    try:
        if data[0] != {}:
            # Получаем ключи и количество колонок
            keys = list(data[0].keys())
            num_columns = 15 #6+9
            columns = ['Факт среднее\n кол-во файлов\n в месяц',
                       'Факт  \nкол-во машин',
                       'Факт \nмаксимальное\n кол-во файлов',
                       'Факт разница\n нагрузки',
                       'Факт \nнагрузка в %',
                       'Факт \nнехватка машин',

                       'Факт \nсреднее кол-во\n файлов в месяц',
                       'Кол-во \nновых УЗ',
                       'Среднее кол-во\n файлов новых\n УЗ в месяц',
                       'Среднее  кол-во \nфайлов с учетом \nновых УЗ в месяц',
                       'Факт  \nкол-во машин',
                       'Факт \nмаксимальное \nкол-во файлов',
                       'Планируемая\nразница нагрузки',
                       'Планируемая\nнагрузка в %',
                       'Планируемая\nнехватка машин'
                       ]
            l = 2
            if data[1] == {}:
                num_columns = 6
                columns = columns[:6]
                l = 1
            headers = [""] + columns
            table_data = [headers]

            for key in keys:
                row = [key]
                for i in range(l):
                    values = data[i].get(key, [])
                    row.extend(values)
                table_data.append(row)

            output = BytesIO()
            pdf = SimpleDocTemplate(output, pagesize=landscape(letter), rightMargin=20, leftMargin=20, topMargin=20,
                                    bottomMargin=20)
            table = Table(table_data)

            style = TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, 0), 'Arial'),
                ('FONTSIZE', (0, 0), (-1, -1), 6),
                ('BOTTOMPADDING', (0, 0), (-1, 0), 6),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.black)
            ])
            table.setStyle(style)
            table._argW = [pdf.width / len(table_data[0])] * len(table_data[0])

            pdf.build([table])
            output.seek(0)

            return send_file(output, as_attachment=True, download_name="report.pdf", mimetype="application/pdf")
        else:
            return '', 204
    except Exception:
        return '', 204

#экспорт в вордовский документ
@app.route('/export_word')
def export_word():
    try:
        if data[0] != {}:

            doc = Document()

            # Добавляем текст с ключевыми словами для замены
            doc.add_paragraph(f"На данный момент в отделе работает 15 машин, из них 8 работают в сменном графике (2/2) по 12 часов с 08:00 до 20:00 и с 20:00 до 08:00 и 7 работают по пятидневной рабочей неделе в промежутке времени с 08:00 до 20:00, что позволяет нам обрабатывать {round(max_files_month['180h'] + max_files_month['168h'] + max_files_month['79h'] + max_files_month['180h_pr'] + max_files_month['180h_night'])} файлов в месяц.\n")
            doc.add_paragraph("\tТекущие показатели эффективности работы отдела:\n")
            doc.add_paragraph(f"•\tСреднее время обработки одного файла в дневное время — {time_obr_day} минут.")
            doc.add_paragraph(f"•\tСреднее время обработки одного файла в ночное время и выходные дни — {time_obr_night} минут.")
            doc.add_paragraph("•\tФактическое количество новых пользователей — 8300.")
            doc.add_paragraph(f"•\tФактическое количество файлов в дневное время — {data[0]['180h'][2] + data[0]['168h'][2] + data[0]['79h'][2]}.")
            doc.add_paragraph(f"•\tФактическое количество файлов в ночное время — {data[0]['180h_night'][2]}.")
            doc.add_paragraph("•\tФактическое количество машин:")
            doc.add_paragraph(f"\t‣\t180ч — {data[0]['180h'][1]};")
            doc.add_paragraph(f"\t‣\t168ч — {data[0]['168h'][1]};")
            doc.add_paragraph(f"\t‣\t79ч — {data[0]['79h'][1]};")
            doc.add_paragraph(f"\t‣\t180ч ночь — {data[0]['180h_night'][1]};")
            doc.add_paragraph(f"•\tПроцент фактической нагрузки на одну машину составляет — ")
            doc.add_paragraph(f"\t‣\t180-79ч — {data[0]['180h'][4]}%;")
            doc.add_paragraph(f"\t‣\t180ч праздники — {data[0]['180h_pr'][4]}%;")
            doc.add_paragraph(f"\t‣\t180ч ночь — {data[0]['180h_night'][4]}%")
            doc.add_paragraph("•\tФактическое количество нехватки машин:")
            doc.add_paragraph(f"\t‣\t180ч — {data[0]['180h'][5]};" )
            doc.add_paragraph(f"\t‣\t168ч — {data[0]['168h'][5]};" )
            doc.add_paragraph(f"\t‣\t79ч — {data[0]['79h'][5]};" )
            doc.add_paragraph(f"\t‣\t180ч ночь/пр — {data[0]['180h_night'][5]}.\n")

            if data[1] != {}:
                doc.add_paragraph("\n\tПланируемые показатели эффективности работы отдела:\n")
                doc.add_paragraph(f"•\tПланируемое количество новых пользователей — {data[1]['180h'][1]}.")
                doc.add_paragraph(f"•\tПланируемое количество файлов в дневное время — {data[1]['180h'][2]}.")
                doc.add_paragraph(f"•\tПланируемое количество файлов в ночное время — {data[1]['180h_night'][2]}.")
                doc.add_paragraph(f"•\tПланируемое количество новых пользователей с учетом новых пользователей — 15000.")
                doc.add_paragraph(f"•\tПланируемое количество файлов с учетом новых пользователей в дневное время — {data[1]['180h'][3]}.")
                doc.add_paragraph(f"•\tПланируемое количество файлов с учетом новых пользователей в ночное время — {data[1]['180h_night'][3]}.")
                doc.add_paragraph(f"•\tПроцент планируемой нагрузки на одну машину составляет — ")
                doc.add_paragraph(f"\t‣\t180-79ч — {data[1]['180h'][7]}%;")
                doc.add_paragraph(f"\t‣\t180ч праздники — {data[1]['180h_pr'][7]}%")
                doc.add_paragraph(f"\t‣\t180ч ночь — {data[1]['180h_night'][7]}%.")
                doc.add_paragraph(f"•\tПланируемое количество нехватки машин:")
                doc.add_paragraph(f"\t‣\t180ч — {data[1]['180h'][8]};")
                doc.add_paragraph(f"\t‣\t168ч — {data[1]['168h'][8]};")
                doc.add_paragraph(f"\t‣\t79ч — {data[1]['79h'][8]};")
                doc.add_paragraph(f"\t‣\t180ч ночь/пр — {data[1]['180h_night'][8]}.")


            # Сохраняем документ в памяти
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            # Возвращаем файл на скачивание
            return send_file(file_stream, as_attachment=True, download_name="report.docx")
        else:
            return '', 204
    except Exception as ex:
        print(ex)
        return '', 204


#эта функция срабатывает, когда нажали кнопку рассчитать(в index.html)
@app.route('/calculate', methods=['POST'])
def calculate():
    #очищаем словари с данными, полученными после предыдущего рассчета
    data[0].clear()
    data[1].clear()

    ###########################
    # Получаем данные из формы:
    ###########################
    try:
        #факт среднее количество файлов в месяц
        total_files = int(request.form['total_files']) #общее
        day_files = int(request.form['day_files']) #день
        night_files = int(request.form['night_files']) #ночь/пр/вых
        day_pr_files = int(request.form['day_pr_files']) #день/пр/вых

        # факт количетсво машин
        machines_180h = int(request.form['machines_180h'])  # 180 часов
        machines_168h = int(request.form['machines_168h'])  # 168 часов
        machines_79h = int(request.form['machines_79h'])  # 79 часов
        machines_180h_night = int(request.form['machines_180h_night'])  # 180 часов ночь
    except:
        return jsonify(result= ("вы ввели не все данные"))

    isNewUsersexist = False
    #Кол-во новых пользователей (УЗ)
    try:
        new_users = int(request.form['new_users']) #общее
        isNewUsersexist = True
    except:
        pass


    ###############
    #Расчеты:
    ###############

    ###############
    #расчет факт
    ###############
    fact_result = fact(total_files, day_files, night_files, day_pr_files, machines_180h, machines_168h, machines_79h, machines_180h_night)

    ###############
    #рассчет план
    ###############
    if isNewUsersexist:
        plan_result = plan(total_files, day_files, night_files, day_pr_files, machines_180h, machines_168h, machines_79h,
                       machines_180h_night, new_users)
    else:
        plan_result = ('')


    ###########################################
    # Создаём итоговый результат в виде строк
    ###########################################
    result = fact_result + plan_result
    #возвращаем полученные расчеты (они отобразяться в Итоговые результаты id=result-output)
    return jsonify(result=result)


if __name__ == '__main__':
    app.run(debug=True)
